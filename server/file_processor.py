"""File processing pipeline — PDF, DOCX, images, and Google Sheets parsing."""

import base64
import io
import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

logger = logging.getLogger("paper-assistant.file-processor")


# ---------------------------------------------------------------------------
# Data model
# ---------------------------------------------------------------------------


@dataclass
class Section:
    """A named section of the paper (Intro, Methods, Results, etc.)."""

    heading: str
    content: str
    start_line: int = 0
    figures: list[str] = field(default_factory=list)  # figure captions found
    references: list[str] = field(default_factory=list)  # references cited


@dataclass
class FigureInfo:
    """Metadata about an extracted figure."""

    filename: str
    caption: str = ""
    page_number: int = 0
    section: str = ""  # which section it belongs to


@dataclass
class PaperDocument:
    """Parsed paper with sections, figures, and metadata."""

    title: str = ""
    authors: str = ""
    abstract: str = ""
    sections: list[Section] = field(default_factory=list)
    figures: list[FigureInfo] = field(default_factory=list)
    full_text: str = ""
    raw_format: str = "unknown"  # "pdf", "docx", "gdoc", "text"

    # PDF parsing provenance (Tier 0/1 ladder). These are written by parse_pdf
    # and currently read only for status/hint display — no downstream consumer
    # depends on them, so older callers that never set them are unaffected.
    parse_mode: str = "fast"  # "fast" | "auto" | "deep"
    ocr_pages: list[int] = field(default_factory=list)  # 1-indexed pages OCR'd
    ocr_deficient_pages: list[int] = field(default_factory=list)  # pages that needed OCR but didn't get it
    # Why ocr_deficient_pages is non-empty ("" when empty). Lets the UI
    # distinguish "install OCR" from "OCR ran but couldn't recover these":
    #   "not_installed" — OCR deps missing; install to recover
    #   "over_cap"      — OCR installed but deficient pages exceed PDF_OCR_MAX_PAGES
    #   "page_failed"   — OCR ran on the page but returned nothing
    ocr_deficient_reason: str = ""
    figure_ocr_pages: list[int] = field(default_factory=list)  # pages whose embedded images were OCR'd
    figure_ocr_count: int = 0  # number of embedded images OCR'd


@dataclass
class ReviewerComment:
    """A single reviewer comment extracted from feedback."""

    id: str  # e.g., "R1-C3"
    reviewer: str  # "Reviewer 1" or "Editor"
    comment_number: int
    text: str
    severity: str = "unspecified"  # "major", "minor", "editorial"
    related_sections: list[str] = field(default_factory=list)
    related_figures: list[str] = field(default_factory=list)


# ---------------------------------------------------------------------------
# PDF processing
# ---------------------------------------------------------------------------

# One parse cache for Phase 1 (decision 6): keyed by (content hash, mode,
# parser version). Sits at the lowest parse level so both the manuscript path
# (drive_sync) and the named-file scan path (chat_handler) benefit. Capped so a
# long session doesn't accumulate every PDF ever opened.
from collections import OrderedDict as _OrderedDict

_PDF_PARSE_CACHE: "_OrderedDict[tuple[str, str, int], PaperDocument]" = _OrderedDict()
_PDF_PARSE_CACHE_MAX = 16


def _cache_key(content: bytes, mode: str, figure_ocr: bool, text_layer_override: str = "") -> tuple[str, str, int, int, str]:
    import hashlib

    from config import PDF_PARSER_VERSION

    override_hash = (
        hashlib.sha256(text_layer_override.encode("utf-8")).hexdigest()[:16]
        if text_layer_override
        else ""
    )
    return (hashlib.sha256(content).hexdigest()[:16], mode, int(figure_ocr), PDF_PARSER_VERSION, override_hash)


def _native_text_is_deficient(text: str) -> bool:
    """True when a page's native text layer is too thin/garbled to be useful.

    Covers the cases OCR is meant to rescue: empty text (image-only / scanned
    pages), very short text (figure pages with a stray caption fragment), and
    garbled text (corrupted/CID-mapped fonts that decode to replacement chars).
    """
    from config import PDF_OCR_MIN_NATIVE_CHARS

    if not text:
        return True
    stripped = text.strip()
    if len(stripped) < PDF_OCR_MIN_NATIVE_CHARS:
        return True
    # Garbled: high ratio of null / replacement / control chars.
    garbled = sum(
        1 for c in text
        if c in ('\x00', '�') or (ord(c) < 9 and c not in '\n\r\t')
    )
    return (garbled / max(len(text), 1)) > 0.3


def _ocr_pages(content: bytes, page_indices_0: list[int]) -> dict[int, str]:
    """OCR the given 0-indexed pages, returning {page_index_0: text or ""}.

    Renders each page with PyMuPDF at PDF_OCR_RENDER_DPI and runs RapidOCR.
    Any import/init/render/OCR failure yields "" for that page (caller treats
    it as "not recovered"). All heavy imports are local so a missing OCR group
    never breaks Tier 0.
    """
    from config import PDF_OCR_RENDER_DPI

    out: dict[int, str] = {idx: "" for idx in page_indices_0}
    try:
        import fitz  # type: ignore
        import numpy as np  # type: ignore

        from pdf_capabilities import get_ocr_engine

        engine = get_ocr_engine()
        if engine is None:
            return out
        with fitz.open(stream=content, filetype="pdf") as d:
            for idx in page_indices_0:
                try:
                    if idx >= d.page_count:
                        continue
                    pix = d[idx].get_pixmap(dpi=PDF_OCR_RENDER_DPI, alpha=False)
                    img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(
                        pix.height, pix.width, pix.n
                    )
                    result = engine(img)
                    txts = getattr(result, "txts", None) or ()
                    out[idx] = "\n".join(t for t in txts if t).strip()
                except Exception as exc:
                    logger.warning("OCR failed for page %d: %s", idx + 1, exc)
                    out[idx] = ""
    except Exception as exc:
        # Missing optional deps (fitz / numpy / rapidocr) — leave all "".
        logger.info("OCR unavailable for this parse: %s", exc)
    return out


def _ocr_embedded_images(content: bytes, skip_pages: set[int]) -> list[tuple[int, str]]:
    """OCR text inside embedded figure images, returning [(page_num_1, text), ...].

    Unlike :func:`_ocr_pages` (which renders whole deficient pages), this
    targets individual embedded raster images on *any* page — so a figure
    (chart labels, a screenshot of text, a diagram) sitting on a page that
    also has body text still gets its text recovered. Guarded to stay cheap and
    quiet on figure-heavy docs:

    - Dedupes by PyMuPDF xref (one image reused on many pages is OCR'd once).
    - Skips images below ``PDF_OCR_IMAGE_MIN_PIXELS`` (logos / icons / scraps).
    - Skips images on pages in ``skip_pages`` (already page-level-OCR'd).
    - Keeps an image's text only if it has ≥ ``PDF_OCR_FIGURE_MIN_CHARS``
      alphanumeric chars (drops blank / decorative / noisy images).
    - Stops after ``PDF_OCR_MAX_IMAGES`` images (caps worst-case time).

    All heavy imports are local so a missing OCR group never breaks Tier 0.
    """
    from config import (
        PDF_OCR_EMBEDDED_IMAGES,
        PDF_OCR_FIGURE_MIN_CHARS,
        PDF_OCR_IMAGE_MIN_PIXELS,
        PDF_OCR_MAX_IMAGES,
    )

    results: list[tuple[int, str]] = []
    if not PDF_OCR_EMBEDDED_IMAGES:
        return results
    try:
        import fitz  # type: ignore

        from pdf_capabilities import get_ocr_engine

        engine = get_ocr_engine()
        if engine is None:
            return results
        seen_xrefs: set[int] = set()
        kept = 0
        with fitz.open(stream=content, filetype="pdf") as d:
            for pno in range(d.page_count):
                page_num = pno + 1  # 1-indexed
                is_skip_page = page_num in skip_pages
                try:
                    images = d[pno].get_images(full=True)
                except Exception:
                    images = []
                for img_info in images:
                    xref = img_info[0] if img_info else 0
                    if not xref or xref in seen_xrefs:
                        continue
                    seen_xrefs.add(xref)
                    if is_skip_page:
                        # Page already rendered+OCR'd as deficient — its image
                        # pixels were read there; don't duplicate.
                        continue
                    if kept >= PDF_OCR_MAX_IMAGES:
                        return results
                    try:
                        ex = d.extract_image(xref)
                        raw = ex.get("image", b"")
                        w = ex.get("width", 0) or 0
                        h = ex.get("height", 0) or 0
                    except Exception as exc:
                        logger.debug("figure-OCR: extract failed xref=%s: %s", xref, exc)
                        continue
                    if not raw or (w * h) < PDF_OCR_IMAGE_MIN_PIXELS:
                        continue
                    try:
                        result = engine(raw)
                        txts = getattr(result, "txts", None) or ()
                        text = " ".join(t for t in txts if t).strip()
                    except Exception as exc:
                        logger.debug("figure-OCR: ocr failed xref=%s: %s", xref, exc)
                        continue
                    # Keep only meaningful text (filters blank/noisy images).
                    alnum = sum(1 for c in text if c.isalnum())
                    if alnum < PDF_OCR_FIGURE_MIN_CHARS:
                        continue
                    results.append((page_num, text))
                    kept += 1
    except Exception as exc:
        logger.info("figure-OCR unavailable for this parse: %s", exc)
    return results


def parse_pdf(content: bytes, filename: str, mode: str = "auto", figure_ocr: bool = True, text_layer_override: str = "") -> PaperDocument:
    """Extract text and figures from a PDF.

    mode:
      "fast" — Tier 0: pdfplumber native text layer only (always available).
      "auto" — Tier 1: Fast + page-level OCR (RapidOCR + PyMuPDF) on pages whose
               native text is empty/short/garbled. Degrades to Fast when the
               optional OCR deps are missing; pages it could not read are listed
               in ``doc.ocr_deficient_pages`` so the UI can hint the install.
    figure_ocr:
      When True (default), recover text *inside* embedded figure images via
      ``_ocr_embedded_images`` and emit ``[Figure text (page N, OCR): …]``
      blocks. When False, skip that pass — used at Load Project for speed
      (figure OCR is the slow part); page-level OCR on deficient pages still
      runs. Scan-and-keep passes True for the full treatment.

    text_layer_override:
      When non-empty, replaces the reconstructed per-page text with this
      authoritative body text. Used for Google Docs, whose PDF export
      fragments text around comment anchors (a "cut" at the start of a
      commented range) while the text/plain export keeps the body contiguous.
      The PDF is still parsed for figures and per-image figure OCR, so those
      blocks are preserved and appended as usual; page-level OCR is skipped
      since the override already provides the body text.
    """
    import pdfplumber

    from config import (
        PDF_DEFAULT_MODE,
        PDF_OCR_EMBEDDED_IMAGES,
        PDF_OCR_ENABLED,
        PDF_OCR_MAX_PAGES,
    )

    # Resolve the effective mode: "auto" only stays auto if the master switch
    # is on; everything else falls back to Fast.
    if mode not in ("fast", "auto", "deep"):
        mode = PDF_DEFAULT_MODE
    if mode == "deep":
        # Tier 2 deferred — degrade to auto.
        mode = "auto"
    if mode == "auto" and not PDF_OCR_ENABLED:
        mode = "fast"

    # Authoritative body text supplied by the caller (e.g. a Google Doc's
    # text/plain export). Empty/whitespace → no override.
    override = (text_layer_override or "").strip()

    # Parse cache (decision 6).
    key = _cache_key(content, mode, figure_ocr, override)
    cached = _PDF_PARSE_CACHE.get(key)
    if cached is not None:
        _PDF_PARSE_CACHE.move_to_end(key)
        logger.info("Parsed PDF '%s' (mode=%s): cache hit", filename, mode)
        return cached

    doc = PaperDocument(title=filename, raw_format="pdf", parse_mode=mode)
    full_text_parts: list[str] = []
    all_figures: list[FigureInfo] = []

    # Pass 1: native text + deficiency flag + image metadata, per page.
    page_native: list[tuple[int, str, bool]] = []  # (page_num_1, text, deficient)
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            deficient = _native_text_is_deficient(text)
            page_native.append((i, text, deficient))
            if hasattr(page, "images") and page.images:
                for img_idx, _img in enumerate(page.images):
                    all_figures.append(
                        FigureInfo(
                            filename=f"page{i}_img{img_idx}.png",
                            page_number=i,
                        )
                    )

    deficient_pages = [p for p, _, d in page_native if d]

    # Decide whether OCR runs at all.
    ocr_pages_run: list[int] = []          # 1-indexed pages actually OCR'd
    ocr_deficient: list[int] = []          # 1-indexed pages left unrecovered
    ocr_deficient_reason: str = ""         # "not_installed" | "over_cap" | "page_failed"
    ocr_text_by_page: dict[int, str] = {}  # 0-indexed -> recovered text

    if mode == "auto" and deficient_pages and not override:
        from pdf_capabilities import get_pdf_capabilities

        caps = get_pdf_capabilities()
        over_cap = len(deficient_pages) > PDF_OCR_MAX_PAGES
        if not caps["auto"]:
            # OCR deps missing — flag every deficient page for the UI hint.
            ocr_deficient = list(deficient_pages)
            ocr_deficient_reason = "not_installed"
        elif over_cap:
            logger.info(
                "PDF '%s': %d deficient pages exceed OCR cap (%d); skipping OCR",
                filename, len(deficient_pages), PDF_OCR_MAX_PAGES,
            )
            ocr_deficient = list(deficient_pages)
            ocr_deficient_reason = "over_cap"
        else:
            recovered = _ocr_pages(content, [p - 1 for p in deficient_pages])
            for p in deficient_pages:
                text = recovered.get(p - 1, "")
                if text:
                    ocr_pages_run.append(p)
                    ocr_text_by_page[p] = text
                else:
                    ocr_deficient.append(p)
            if ocr_deficient:
                ocr_deficient_reason = "page_failed"

    # Pass 2: assemble full_text in document order, substituting OCR text for
    # recovered pages and keeping native text (even if empty) otherwise. When a
    # caller supplied an authoritative text layer, use it as the sole body text
    # instead of the reconstructed per-page text (which, for Google Docs, is
    # fragmented around comment anchors).
    if override:
        full_text_parts.append(override)
    else:
        for i, native_text, _deficient in page_native:
            if i in ocr_text_by_page:
                full_text_parts.append(ocr_text_by_page[i])
            else:
                full_text_parts.append(native_text)

    # Per-image figure OCR (auto mode only): recover text inside embedded
    # figure images on any page. Skips pages already page-level-OCR'd above.
    # Gated on ``figure_ocr`` — Load Project passes False to skip this slow
    # pass; scan-and-keep passes True for the full treatment.
    figure_ocr_pages: list[int] = []
    if mode == "auto" and PDF_OCR_EMBEDDED_IMAGES and figure_ocr:
        figure_hits = _ocr_embedded_images(content, skip_pages=set(ocr_pages_run))
        for page_num, text in figure_hits:
            full_text_parts.append(f"[Figure text (page {page_num}, OCR): {text}]")
            figure_ocr_pages.append(page_num)

    doc.full_text = "\n\n".join(full_text_parts)
    doc.figures = all_figures
    doc.ocr_pages = ocr_pages_run
    doc.ocr_deficient_pages = ocr_deficient
    doc.ocr_deficient_reason = ocr_deficient_reason
    doc.figure_ocr_pages = figure_ocr_pages
    doc.figure_ocr_count = len(figure_ocr_pages)
    doc.sections = _parse_sections(doc.full_text)
    doc.title, doc.abstract, doc.authors = _extract_metadata(doc.full_text)

    # Store in cache.
    _PDF_PARSE_CACHE[key] = doc
    _PDF_PARSE_CACHE.move_to_end(key)
    while len(_PDF_PARSE_CACHE) > _PDF_PARSE_CACHE_MAX:
        _PDF_PARSE_CACHE.popitem(last=False)

    logger.info(
        "Parsed PDF '%s' (mode=%s): %d pages, %d sections, %d images, "
        "%d page-OCR'd, %d deficient-unread, %d figure-images-OCR'd",
        filename, mode, len(page_native), len(doc.sections), len(all_figures),
        len(ocr_pages_run), len(ocr_deficient), len(figure_ocr_pages),
    )
    return doc


# ---------------------------------------------------------------------------
# DOCX processing
# ---------------------------------------------------------------------------

# A figure-caption paragraph: an optional prefix (Supplemental/Supplementary/
# Suppl./Supp./Extended Data), then "Figure"/"Fig.", a label token, and a
# delimiter (period/colon/dash/paren). The label token is broad — "1", "S1",
# "S18ex", "1A" — so any "Figure <label>." style is counted; the figure NUMBER
# inside the label (extracted separately) is what we use for ordering and
# same-figure dedup, not the label string itself. The delimiter weeds out
# mid-sentence references such as "Figure 1 shows..." so we count real
# captions, not in-text mentions. DOCX has no reliable embedded-image→caption
# link, so we key figures by their caption label instead of by image blob.
# Group 1 = prefix (for family detection), group 2 = label token.
_FIGURE_CAPTION_RE = re.compile(
    r"^\s*"
    r"((?:Supplemental|Supplementary|Suppl\.?|Supp\.?|Extended\s+Data)\s+)?"
    r"(?:Figure|Fig\.?)\s+"
    r"([A-Za-z]?\d+[A-Za-z0-9]*)\s*"
    r"[.:)\-–—]",
    re.IGNORECASE,
)


def _figure_number(label_token: str) -> Optional[int]:
    """Extract the leading integer from a figure label token.

    "S18ex" → 18, "1A" → 1, "S1" → 1, "20" → 20. Used to order figures and to
    spot that "Figure S1" (index) and "Figure 1" (body) are the same figure.
    """
    m = re.search(r"\d+", label_token)
    return int(m.group()) if m else None


def _figure_family(prefix: Optional[str]) -> str:
    """Coarse caption family: "supp" (Supplemental/Supplementary), "ext"
    (Extended Data), or "main" (no prefix). Two captions with the same figure
    NUMBER but different families are NOT the same figure — e.g. a main
    "Figure 1" vs a "Supplementary Figure S1" — so family guards the dedup."""
    if not prefix:
        return "main"
    p = prefix.lower()
    if "extended" in p:
        return "ext"
    return "supp"


def _dedupe_figures(raw: list[dict]) -> list[FigureInfo]:
    """Collapse duplicate figure captions and order survivors by figure number.

    Supplemental files commonly open with an index that re-lists every figure
    by title, then repeat each caption in the body (sometimes with its full
    legend). The two copies often use different labels AND different title
    wording — "Supplemental Figure S2. Hi-C maps of scaffolds." in the index
    vs "Supplemental Figure 2. Hi-C contact maps of the assembly…" in the body
    — so neither label nor caption-text similarity reliably connects them. The
    one stable signal is the figure NUMBER, so we treat two captions as the
    SAME figure when they share the same number AND family (both Supplemental,
    both Extended Data, or both plain main-text figures), and keep the longest
    caption (the legend-bearing body copy). Family prevents merging a main
    "Figure 1" with a "Supplementary Figure S1" that happens to share number 1.
    Survivors are ordered by figure number so the list follows figure order
    regardless of how labels were formatted.
    """
    kept: list[dict] = []
    for e in raw:
        dup_of = None
        for k in kept:
            if (k["num"] is not None and k["family"] == e["family"]
                    and k["num"] == e["num"]):
                dup_of = k
                break
        if dup_of is None:
            kept.append(e)
        elif len(e["desc"]) > len(dup_of["desc"]):
            # Replace with the richer caption but preserve the earlier
            # position so final ordering stays stable.
            e["idx"] = dup_of["idx"]
            kept[kept.index(dup_of)] = e
        # else: a shorter duplicate of an already-kept caption → drop.
    kept.sort(key=lambda e: (e["num"] if e["num"] is not None else float("inf"), e["idx"]))
    return [FigureInfo(filename=f"Figure {e['label']}", caption=e["caption"]) for e in kept]


# DOCX figure-OCR — qualified names for scraping embedded-image rIds from a
# paragraph's XML. python-docx oxml elements don't accept .xpath(namespaces=),
# so we use findall with qualified names. ``qn`` has no ``v:`` prefix (VML isn't
# in its nsmap), so VML imagedata uses the raw qualified name.
_R_EMBED = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
_R_ID = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"
_VML_IMAGEDATA = "{urn:schemas-microsoft-com:vml}imagedata"
# WordprocessingML block containers we descend into during docx parsing.
# Structured Document Tags (w:sdt) wrap content controls — some Word
# manuscript templates put figure tables inside one — and w:sdtContent holds
# the actual block children. Without descending, everything inside an SDT
# (figures, captions, legends) is silently dropped from full_text.
_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_SDT = _W + "sdt"
_SDT_CONTENT = _W + "sdtContent"


def _paragraph_image_rids(paragraph) -> list[str]:
    """rIds of images embedded in a paragraph (DrawingML blips + legacy VML)."""
    try:
        from docx.oxml.ns import qn

        blips = paragraph._p.findall(".//" + qn("a:blip"))
        vmls = paragraph._p.findall(".//" + _VML_IMAGEDATA)
        rids = [b.get(_R_EMBED) for b in blips] + [v.get(_R_ID) for v in vmls]
        return [r for r in rids if r]
    except Exception:
        return []


def _ocr_docx_images(docx_part) -> dict:
    """OCR text inside embedded images of a .docx, returning {rId: text}.

    Mirrors :func:`_ocr_embedded_images` (the PDF figure-OCR path) but for
    Word documents: targets each image relationship on the main document part
    instead of PDF xrefs. Only runs when the optional OCR group is installed
    (``get_ocr_engine()`` returns a usable engine) and the
    ``PDF_OCR_EMBEDDED_IMAGES`` master switch is on — otherwise returns ``{}``
    so ``parse_docx`` degrades to today's captions-only behavior (no crash).

    Guards (same as the PDF path): dedupes by image ``partname`` (one image
    reused via many rIds is OCR'd once, all its rIds mapped to the text);
    skips images below ``PDF_OCR_IMAGE_MIN_PIXELS`` and images whose format
    can't be read (EMF/WMF vector metafiles raise on ``part.image``); keeps an
    image's text only at ≥ ``PDF_OCR_FIGURE_MIN_CHARS`` alphanumeric chars; and
    stops after ``PDF_OCR_MAX_IMAGES``.
    """
    from config import (
        PDF_OCR_EMBEDDED_IMAGES,
        PDF_OCR_FIGURE_MIN_CHARS,
        PDF_OCR_IMAGE_MIN_PIXELS,
        PDF_OCR_MAX_IMAGES,
    )

    rid_to_text: dict[str, str] = {}
    if not PDF_OCR_EMBEDDED_IMAGES:
        return rid_to_text
    try:
        from pdf_capabilities import get_ocr_engine

        engine = get_ocr_engine()
    except Exception:
        return rid_to_text
    if engine is None:
        return rid_to_text

    try:
        rels = docx_part.rels
    except Exception:
        return rid_to_text

    # Group rIds by the underlying image part (dedupe by partname).
    part_to_rids: dict = {}
    for rid, rel in rels.items():
        if not rel.reltype.endswith("/image"):
            continue
        try:
            part = rel.target_part
        except Exception:
            continue
        key = str(getattr(part, "partname", "")) or id(part)
        bucket = part_to_rids.get(key)
        if bucket is None:
            part_to_rids[key] = (part, [rid])
        else:
            bucket[1].append(rid)

    kept = 0
    for part, rids in part_to_rids.values():
        if kept >= PDF_OCR_MAX_IMAGES:
            break
        # Dimensions for the min-pixels guard. Accessing .image raises for
        # unsupported formats (EMF/WMF) — skip those (can't OCR vector data).
        try:
            img = part.image
            w, h = img.px_width, img.px_height
        except Exception:
            continue
        if (w * h) < PDF_OCR_IMAGE_MIN_PIXELS:
            continue
        try:
            result = engine(part.blob)
            txts = getattr(result, "txts", None) or ()
            text = " ".join(t for t in txts if t).strip()
        except Exception:
            continue
        if sum(1 for c in text if c.isalnum()) < PDF_OCR_FIGURE_MIN_CHARS:
            continue
        for rid in rids:
            rid_to_text[rid] = text
        kept += 1
    return rid_to_text


def parse_docx(content: bytes, filename: str, figure_ocr: bool = True) -> PaperDocument:
    """Extract text from a .docx file, including table cell text and figure OCR.

    python-docx exposes ``.paragraphs`` and ``.tables`` as separate lists,
    which discards table contents *and* loses document order. We walk the
    body's block-level children (``w:p`` / ``w:tbl``) in order so tables are
    rendered where they actually sit in the document, alongside their
    captions and surrounding paragraphs.

    When the optional OCR group is installed, text inside embedded figure
    images is recovered via :func:`_ocr_docx_images` and emitted inline as
    ``[Figure text (OCR): …]`` blocks at each figure's position — so a Word
    supplement's figures get their axis labels / titles / table headers read,
    not just their captions. Without OCR, degrades to captions-only (no crash).

    ``figure_ocr`` (default True) gates that per-image figure-OCR pass. Load
    Project passes False for speed (figure OCR is the slow part); text,
    captions, and SDT/table content are still extracted. Scan-and-keep passes
    True for the full treatment.
    """
    from docx import Document as DocxDocument
    from docx.document import Document as _Document
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table, _Cell
    from docx.text.paragraph import Paragraph

    doc = PaperDocument(title=filename, raw_format="docx")
    docx = DocxDocument(io.BytesIO(content))

    # {rId: ocr_text} for embedded figure images. Empty when OCR is unavailable
    # (not installed / master switch off) or when the caller opted out via
    # figure_ocr=False (Load Project) — the inline emit below is then a no-op.
    rid_to_text = _ocr_docx_images(docx.part) if figure_ocr else {}
    figure_ocr_count = 0

    def iter_block_items(parent):
        """Yield Paragraph and Table objects in document order.

        Descends into Structured Document Tags (w:sdt → w:sdtContent) so
        content controls — some Word manuscript templates wrap figure tables
        in one — are not silently dropped. Recurses through nested SDTs.
        """
        if isinstance(parent, _Document):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:
            return

        def _walk(elm):
            for child in elm.iterchildren():
                if isinstance(child, CT_P):
                    yield Paragraph(child, parent)
                elif isinstance(child, CT_Tbl):
                    yield Table(child, parent)
                elif child.tag == _SDT:
                    content = child.find(_SDT_CONTENT)
                    if content is not None:
                        yield from _walk(content)

        yield from _walk(parent_elm)

    def process_paragraph(p: Paragraph) -> tuple[list[str], dict | None, int]:
        """Emit text + inline figure-OCR for one paragraph; detect captions.

        Shared by the body-paragraph loop and table-cell rendering so figures
        embedded inside table cells (common when a manuscript wraps figure
        blocks in a table) get the same OCR-emit + caption-detect treatment as
        body-level figures. Returns ``(parts_to_emit, raw_figure_or_None,
        n_ocr_blocks)``. Whitespace-only paragraphs emit nothing except any
        recovered figure-OCR text.
        """
        parts: list[str] = []
        n_ocr = 0
        # Figure-image OCR: emit recovered text inline at the image's
        # paragraph so it sits adjacent to the caption, not lost entirely.
        if rid_to_text:
            for rid in _paragraph_image_rids(p):
                fig_text = rid_to_text.get(rid)
                if fig_text:
                    parts.append(f"[Figure text (OCR): {fig_text}]")
                    n_ocr += 1
        text = p.text
        if text and text.strip():
            parts.append(text)
        # Detect figure-caption paragraphs. The label token (e.g. "1",
        # "S1", "S18ex", "1A") carries the figure number used for ordering
        # and same-figure dedup; the prefix drives the caption family so a
        # main "Figure 1" is never merged with a "Supplementary Figure S1".
        raw_fig: dict | None = None
        m = _FIGURE_CAPTION_RE.match(text)
        if m:
            prefix, label_token = m.group(1), m.group(2)
            raw_fig = {
                "label": label_token,
                "desc": text[m.end():].strip().lower(),
                "num": _figure_number(label_token),
                "family": _figure_family(prefix),
                "idx": len(raw_figures),
                "caption": text.strip(),
            }
        return parts, raw_fig, n_ocr

    def table_to_text(table: Table) -> tuple[str, int, int]:
        """Render a table as pipe-delimited rows (markdown-style).

        Returns ``(text, n_paragraphs, n_figure_ocr)``. Cell paragraphs are
        run through :func:`process_paragraph` so figure images embedded in
        cells emit inline ``[Figure text (OCR): …]`` and their captions are
        collected — figures wrapped in a layout table are not lost.
        """
        rows = []
        n_paras = 0
        n_ocr = 0
        for row in table.rows:
            cells = []
            for cell in row.cells:
                # A cell may hold multiple paragraphs; flatten each to its
                # text + any inline figure-OCR, then collapse internal
                # whitespace so each row stays a single line.
                cell_parts: list[str] = []
                for p in cell.paragraphs:
                    n_paras += 1
                    parts, raw_fig, ocr = process_paragraph(p)
                    if raw_fig:
                        raw_figures.append(raw_fig)
                    n_ocr += ocr
                    cell_parts.extend(parts)
                cell_text = re.sub(r"\s+", " ", "\n".join(cell_parts)).strip()
                cells.append(cell_text)
            rows.append(" | ".join(cells))
        if not rows:
            return "", n_paras, n_ocr
        # First row as header + separator → a markdown table the LLM reads
        # as structured data rather than free text.
        header = rows[0]
        sep = " | ".join("---" for _ in rows[0].split(" | "))
        return "\n".join([header, sep, *rows[1:]]), n_paras, n_ocr

    full_text_parts = []
    n_paragraphs = 0
    n_tables = 0
    # Collect every figure-caption paragraph; dedup happens after the loop so
    # we can compare an index entry against the body caption of the same figure
    # and keep the richer (legend-bearing) copy.
    raw_figures: list[dict] = []
    for block in iter_block_items(docx):
        if isinstance(block, Paragraph):
            n_paragraphs += 1
            parts, raw_fig, ocr = process_paragraph(block)
            if raw_fig:
                raw_figures.append(raw_fig)
            figure_ocr_count += ocr
            full_text_parts.extend(parts)
        elif isinstance(block, Table):
            rendered, cell_paras, cell_ocr = table_to_text(block)
            n_paragraphs += cell_paras
            figure_ocr_count += cell_ocr
            if rendered:
                full_text_parts.append(rendered)
                n_tables += 1

    # Collapse any residual runs of 3+ newlines (e.g. multi-line table cells
    # that reduced to blank lines) to a single blank line, so the extracted
    # text never presents an image-gap as a wall of empty space.
    doc.full_text = re.sub(r"\n{3,}", "\n\n", "\n\n".join(full_text_parts))
    doc.figures = _dedupe_figures(raw_figures)
    doc.figure_ocr_count = figure_ocr_count
    doc.sections = _parse_sections(doc.full_text)
    doc.title, doc.abstract, doc.authors = _extract_metadata(doc.full_text)

    logger.info("Parsed DOCX '%s': %d paragraphs, %d tables, %d sections, %d figures (%d before dedup), %d figure-images-OCR'd",
                 filename, n_paragraphs, n_tables, len(doc.sections), len(doc.figures), len(raw_figures), figure_ocr_count)
    return doc


# ---------------------------------------------------------------------------
# Markdown / plain text
# ---------------------------------------------------------------------------


def parse_text(content: str, filename: str) -> PaperDocument:
    """Parse plain text or markdown as a paper."""
    doc = PaperDocument(title=filename, raw_format="text")
    doc.full_text = content
    doc.sections = _parse_sections(content)
    doc.title, doc.abstract, doc.authors = _extract_metadata(content)

    logger.info("Parsed text '%s': %d chars, %d sections",
                 filename, len(content), len(doc.sections))
    return doc


# ---------------------------------------------------------------------------
# Section parsing
# ---------------------------------------------------------------------------


def _parse_sections(text: str) -> list[Section]:
    """Split paper text into sections based on common academic headers."""
    from config import SECTION_PATTERNS

    patterns = [re.compile(p, re.IGNORECASE | re.MULTILINE) for p in SECTION_PATTERNS]

    # Find all potential section boundaries
    boundaries: list[tuple[int, str]] = []
    for pattern in patterns:
        for match in pattern.finditer(text):
            boundaries.append((match.start(), match.group().strip()))

    boundaries.sort(key=lambda x: x[0])

    if not boundaries:
        # If no headers detected, treat the entire text as one section
        return [Section(heading="Full Text", content=text)]

    sections = []
    for i, (pos, heading) in enumerate(boundaries):
        start = pos
        end = boundaries[i + 1][0] if i + 1 < len(boundaries) else len(text)
        content = text[start:end].strip()

        # Remove the heading from content start
        if content.startswith(heading):
            content = content[len(heading):].strip()

        sections.append(Section(heading=heading, content=content, start_line=start))

    return sections


# ---------------------------------------------------------------------------
# Metadata extraction
# ---------------------------------------------------------------------------


def _extract_metadata(text: str) -> tuple[str, str, str]:
    """Heuristically extract title, abstract, and authors from the paper text."""
    title = ""
    abstract = ""
    authors = ""

    lines = text.strip().split("\n")

    # Title: first substantial line
    title_idx = -1
    for i, line in enumerate(lines[:20]):
        stripped = line.strip()
        if len(stripped) > 10 and not stripped.startswith(("#", "http", "©", "Correspondence")):
            title = stripped
            title_idx = i
            break

    # Abstract: between "Abstract" and the next section header
    abs_match = re.search(
        r"(?:^|\n)(?:#+\s*)?(?:Abstract|Summary)\s*\n+(.+?)(?:\n(?:#+\s*)?(?:Introduction|Background|Main|Results))",
        text, re.DOTALL | re.IGNORECASE
    )
    if abs_match:
        abstract = abs_match.group(1).strip()[:2000]

    # Authors: scan the few lines after the title for a plausible author list.
    # Conservative — only set when a line looks like comma/and-separated names
    # and not an affiliation or a sentence; otherwise leave empty rather than
    # risk populating garbage.
    if title_idx >= 0:
        affiliation_tokens = ("university", "department", "institute", "laborator",
                              "@", "http", "corresponding", "©", "email")
        for line in lines[title_idx + 1: title_idx + 9]:
            stripped = line.strip()
            if not (10 <= len(stripped) <= 200):
                continue
            low = stripped.lower()
            if stripped.startswith(("#", "http", "©")):
                continue
            if "," not in stripped and not re.search(r"\band\b", low):
                continue
            if stripped.endswith(".") or any(tok in low for tok in affiliation_tokens):
                continue
            # Strip affiliation superscript markers (digits, *, †, ‡, §)
            authors = re.sub(r"[\d\*†‡§]+", "", stripped)
            authors = re.sub(r",\s*,", ",", authors)
            authors = re.sub(r"\s*,\s*", ", ", authors)
            authors = re.sub(r"\s{2,}", " ", authors).strip(" ,")
            break

    return title, abstract, authors


# ---------------------------------------------------------------------------
# Reviewer comment extraction
# ---------------------------------------------------------------------------


def _is_garbled(text: str, threshold: float = 0.3) -> bool:
    """Return True if the text looks like corrupted/binary data."""
    if not text:
        return True
    # Count replacement chars (�), null bytes, and high control chars
    garbled = sum(1 for c in text if c in ('\x00', '�') or (ord(c) < 9 and c not in '\n\r\t'))
    return (garbled / max(len(text), 1)) > threshold


def _deduplicate_comments(comments: list[ReviewerComment]) -> list[ReviewerComment]:
    """Remove comments whose text is substantially similar to another."""
    seen = []
    for c in comments:
        # Compare first 100 chars against already-accepted comments
        prefix = c.text[:100].strip().lower()
        is_dup = any(
            prefix in s.text[:200].strip().lower()
            or s.text[:100].strip().lower() in prefix
            for s in seen
        )
        if not is_dup:
            seen.append(c)
    return seen


def extract_reviewer_comments(text: str) -> list[ReviewerComment]:
    """
    Extract individual reviewer comments from reviewer feedback text.

    Handles common formats:
    - "Reviewer 1, Comment 1: ..."
    - "Reviewer 1 Comments for the Author" + introductory paragraph + numbered points
    - "R1-C1: ..."
    - Numbered lists under reviewer headers
    - Editor / AE comments
    """
    # Pre-filter: strip out obviously garbled chunks
    if _is_garbled(text):
        logger.warning("Skipping entirely garbled comment file")
        return []

    comments = []
    counter = [0]  # mutable counter for unique IDs

    # --- Pattern 1: "Reviewer X, Comment Y: ..." ---
    pattern1 = re.compile(
        r"(?:Reviewer|Referee)\s*(\d+)[,:]\s*(?:Comment|Point|Issue|Question)\s*(\d+)[,:]\s*(.+?)(?=(?:Reviewer|Referee)\s*\d+|$)",
        re.DOTALL | re.IGNORECASE,
    )

    for match in pattern1.finditer(text):
        reviewer_num = match.group(1)
        comment_num = int(match.group(2))
        comment_text = match.group(3).strip()
        if _is_garbled(comment_text) or len(comment_text) < 20:
            continue
        counter[0] += 1
        comments.append(
            ReviewerComment(
                id=f"R{reviewer_num}-C{comment_num}",
                reviewer=f"Reviewer {reviewer_num}",
                comment_number=comment_num,
                text=comment_text[:3000],
                severity=_classify_severity(comment_text),
            )
        )

    # --- Pattern 2: "Reviewer X (Name)" header block with numbered points ---
    pattern2 = re.compile(
        r"(?:Reviewer|Referee)\s*(\d+)\s*(?:[:(].*?[):])?\s*\n(.+?)(?=(?:Reviewer|Referee)\s*\d+|===|Editor\b|AE\b|Associate\s+Editor|$)",
        re.DOTALL | re.IGNORECASE,
    )
    for match in pattern2.finditer(text):
        reviewer_num = match.group(1)
        block = match.group(2).strip()
        if _is_garbled(block):
            continue
        # Extract numbered points
        points = re.split(r"\n\s*(?:\d+[.)]\s*|\*\s*)", block)
        for i, point in enumerate(points, start=1):
            point = point.strip()
            if _is_garbled(point):
                continue
            if len(point) > 20:
                counter[0] += 1
                comments.append(
                    ReviewerComment(
                        id=f"R{reviewer_num}-C{i}",
                        reviewer=f"Reviewer {reviewer_num}",
                        comment_number=i,
                        text=point[:3000],
                        severity=_classify_severity(point),
                    )
                )

    # --- Pattern 3: "Reviewer X Comments for the Author" with intro paragraph + numbers ---
    pattern3 = re.compile(
        r"(?:Reviewer|Referee)\s*(\d+)\s*(?:Comments|Feedback|Report)\s*(?:for\s+the\s+Author[s]?)?[:\n]\s*(.+?)(?=(?:Reviewer|Referee)\s*\d+|Editor\b|AE\b|Associate\s+Editor|===|$)",
        re.DOTALL | re.IGNORECASE,
    )
    for match in pattern3.finditer(text):
        reviewer_num = match.group(1)
        block = match.group(2).strip()
        if _is_garbled(block):
            continue
        # Split into introductory paragraph + numbered points
        parts = re.split(r"\n\s*(?=\d+[.)]\s)", block)
        for part in parts:
            part = part.strip()
            if _is_garbled(part):
                continue
            # Identify if this is a numbered point or the intro paragraph
            num_match = re.match(r"(\d+)[.)]\s*(.+)", part, re.DOTALL)
            if num_match:
                comment_num = int(num_match.group(1))
                point_text = num_match.group(2).strip()
            else:
                # Introductory / general comment
                comment_num = 0
                point_text = part

            if len(point_text) > 20:
                counter[0] += 1
                label = f"C{comment_num}" if comment_num > 0 else "Intro"
                comments.append(
                    ReviewerComment(
                        id=f"R{reviewer_num}-{label}",
                        reviewer=f"Reviewer {reviewer_num}",
                        comment_number=comment_num if comment_num > 0 else counter[0],
                        text=point_text[:3000],
                        severity=_classify_severity(point_text),
                    )
                )

    # --- Pattern 4: Editor / AE comments ---
    editor_pattern = re.compile(
        r"(?:Editor|AE|Associate\s+Editor)\s*(?:\(.*?\))?\s*:?\s*\n?(.+?)(?=(?:Reviewer|Referee)\s*\d+|Editor\b|AE\b|Associate\s+Editor|===|$)",
        re.DOTALL | re.IGNORECASE,
    )
    for match in editor_pattern.finditer(text):
        ed_text = match.group(1).strip()
        if _is_garbled(ed_text):
            continue
        if len(ed_text) > 20:
            counter[0] += 1
            comments.append(
                ReviewerComment(
                    id=f"ED-C{counter[0]}",
                    reviewer="Editor",
                    comment_number=counter[0],
                    text=ed_text[:3000],
                    severity=_classify_severity(ed_text),
                )
            )

    # --- Fallback: free-text chunking ---
    if not comments:
        comments = _extract_from_free_text(text)

    # Deduplicate and log
    comments = _deduplicate_comments(comments)
    logger.info("Extracted %d reviewer comments from text (pre-dedup: %d)", len(comments), counter[0])
    return comments


def extract_reviewer_comments_from_sheets(
    sheets_data: dict,
) -> list[ReviewerComment]:
    """
    Extract reviewer comments from Google Sheets data.

    Expected columns (flexible order, detected by header):
    - Reviewer / Source
    - Comment / Feedback / Concern
    - Severity (optional)
    - Status (optional)
    - Response (optional)
    """
    comments = []
    for sheet_name, rows in sheets_data.items():
        if not rows:
            continue

        # Detect header row
        header = [cell.lower().strip() if cell else "" for cell in rows[0]]

        reviewer_col = _find_column(header, ["reviewer", "source", "from"])
        comment_col = _find_column(header, ["comment", "feedback", "concern", "point", "question"])
        severity_col = _find_column(header, ["severity", "priority", "type", "category"])
        response_col = _find_column(header, ["response", "reply", "answer", "draft"])

        if comment_col is None:
            # Assume first column is reviewer, second is comment
            comment_col = 1
            reviewer_col = 0

        for i, row in enumerate(rows[1:], start=1):
            if not row or len(row) <= (comment_col or 1):
                continue

            comment_text = str(row[comment_col]) if comment_col < len(row) else ""
            if not comment_text.strip() or len(comment_text.strip()) < 10:
                continue

            reviewer_name = (
                str(row[reviewer_col]) if reviewer_col is not None and reviewer_col < len(row)
                else "Unknown"
            )

            severity = "unspecified"
            if severity_col is not None and severity_col < len(row):
                severity = str(row[severity_col]).lower()

            comments.append(
                ReviewerComment(
                    id=f"{sheet_name}-C{i}",
                    reviewer=reviewer_name,
                    comment_number=i,
                    text=comment_text[:3000],
                    severity=severity,
                )
            )

    logger.info("Extracted %d reviewer comments from sheets", len(comments))
    return comments


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _find_column(header: list[str], candidates: list[str]) -> Optional[int]:
    """Find the index of a column matching any candidate name."""
    for idx, cell in enumerate(header):
        for candidate in candidates:
            if candidate in cell:
                return idx
    return None


def _classify_severity(text: str) -> str:
    """Heuristically classify comment severity."""
    text_lower = text.lower()
    major_keywords = ["major", "critical", "must", "essential", "significant flaw",
                       "fatal", "fundamental", "require"]
    minor_keywords = ["minor", "clarify", "suggestion", "could", "optional",
                       "consider", "perhaps", "might"]
    editorial_keywords = ["typo", "grammar", "spelling", "format", "reference",
                           "citation", "punctuation"]

    major_count = sum(1 for kw in major_keywords if kw in text_lower)
    minor_count = sum(1 for kw in minor_keywords if kw in text_lower)
    editorial_count = sum(1 for kw in editorial_keywords if kw in text_lower)

    if major_count > 0:
        return "major"
    elif editorial_count > minor_count:
        return "editorial"
    elif minor_count > 0:
        return "minor"
    return "unspecified"


def _extract_from_free_text(text: str) -> list[ReviewerComment]:
    """Last-resort extraction: split by blank lines or numbered items."""
    comments = []
    chunks = re.split(r"\n\s*\n", text)
    counter = 0
    for chunk in chunks:
        chunk = chunk.strip()
        if len(chunk) > 30:
            counter += 1
            comments.append(
                ReviewerComment(
                    id=f"GEN-C{counter}",
                    reviewer="General",
                    comment_number=counter,
                    text=chunk[:3000],
                    severity=_classify_severity(chunk),
                )
            )
    return comments
